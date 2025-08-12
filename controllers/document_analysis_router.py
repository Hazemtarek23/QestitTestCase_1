from typing import Dict, List, Optional, Any
from fastapi import HTTPException
import logging
import unicodedata
import pathlib
import os
import chardet
import fitz  # PyMuPDF for PDF extraction
import docx  # python-docx for Word documents
from io import BytesIO
from services.llama_client import LlamaClient
from services.ports import TextGenerator


class DocumentChunk:
    def __init__(self, chunk_id: str, content: str):
        self.chunk_id = chunk_id
        self.content = content


class DocumentAnalysisRequest:
    def __init__(self, user_story: str, chunks: List[DocumentChunk]):
        self.user_story = user_story
        self.chunks = chunks


class SingleAnalysisRequest:
    def __init__(self, text: str):
        self.text = text


class TechnicalAnalysisResponse:
    def __init__(self, chunk_id: str, analysis: str):
        self.chunk_id = chunk_id
        self.analysis = analysis


class MultiChunkAnalysisResponse:
    def __init__(self, user_story: str, analyses: List[TechnicalAnalysisResponse]):
        self.user_story = user_story
        self.analyses = analyses


class DocumentAnalysisService:
    def __init__(self, generator: Optional[TextGenerator] = None):
        self.azure_client = generator or LlamaClient()
        self.logger = logging.getLogger(__name__)

        self.TECHNICAL_ANALYSIS_SYSTEM_PROMPT = """You are a technical documentation analysis assistant.         
        Your role is to analyze technical documents and provide comprehensive insights about:
        - System architecture and design patterns
        - Technical requirements and specifications
        - Implementation details and workflows
        - Potential technical risks and considerations
        - Integration points and dependencies

        Provide clear, structured analysis that helps technical teams understand the document's implications."""

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_extension = pathlib.Path(file_path).suffix.lower()

        try:
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_text_from_word(file_path)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']:
                return self._extract_text_from_text_file(file_path)
            else:
                # Try to read as text file with encoding detection
                return self._extract_text_from_text_file(file_path)
        except Exception as e:
            self.logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise ValueError(f"Unable to extract text from file: {str(e)}")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files using PyMuPDF"""
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            if not text.strip():
                raise ValueError("No text content found in PDF")
            return text
        except Exception as e:
            self.logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_text_from_word(self, file_path: str) -> str:
        """Extract text from Word documents using python-docx"""
        try:
            doc = docx.Document(file_path)
            text = ""
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            if not text.strip():
                raise ValueError("No text content found in Word document")
            return text
        except Exception as e:
            self.logger.error(f"Word document extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from Word document: {str(e)}")

    def _extract_text_from_text_file(self, file_path: str) -> str:
        """Extract text from plain text files with encoding detection"""
        try:
            # First, try to detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()

            # Detect encoding
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8')
            confidence = detected.get('confidence', 0)

            self.logger.info(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")

            # Try detected encoding first
            try:
                text = raw_data.decode(encoding)
            except (UnicodeDecodeError, TypeError):
                # Fallback encodings
                encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                text = None
                for enc in encodings_to_try:
                    try:
                        text = raw_data.decode(enc)
                        self.logger.info(f"Successfully decoded with encoding: {enc}")
                        break
                    except UnicodeDecodeError:
                        continue

                if text is None:
                    # Last resort: decode with errors='replace'
                    text = raw_data.decode('utf-8', errors='replace')
                    self.logger.warning("Used UTF-8 with error replacement")

            return text
        except Exception as e:
            self.logger.error(f"Text file extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from file: {str(e)}")

    def _sanitize_text(self, text: str) -> str:
        """Normalize text to avoid encoding issues"""
        if not text:
            return ""

        # Normalize Unicode characters
        text = unicodedata.normalize("NFKC", text)

        # Remove or replace problematic characters
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\ufffd', ' ')  # Replace replacement characters

        # Clean up excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            cleaned_line = ' '.join(line.split())  # Normalize whitespace
            if cleaned_line:  # Skip empty lines
                cleaned_lines.append(cleaned_line)

        return '\n'.join(cleaned_lines)

    async def analyze_file(self, file_path: str) -> str:
        """Analyze any file type (PDF, Word, text) and return a path to analysis output"""
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            # Extract text based on file type
            raw_text = self._extract_text_from_file(file_path)
            if not raw_text.strip():
                raise ValueError(f"No extractable text found in file: {file_path}")

            clean_text = self._sanitize_text(raw_text)
            self.logger.info(f"Analyzing file: {file_path} (length: {len(clean_text)} characters)")

            # Perform analysis using configured LLM
            analysis_result = await self._analyze_document(clean_text)

            # Save analysis result
            output_path = str(pathlib.Path(file_path).with_suffix(".analyzed.txt"))
            with open(output_path, "w", encoding="utf-8") as out_f:
                out_f.write(analysis_result)

            self.logger.info(f"Analysis saved to: {output_path}")
            return output_path

        except FileNotFoundError:
            self.logger.error(f"File not found: {file_path}")
            raise
        except Exception as e:
            self.logger.error(f"Failed to analyze file {file_path}: {str(e)}")
            raise

    async def analyze_technical_document(self, request: SingleAnalysisRequest) -> Dict[str, str]:
        """Technical analysis using the workflow system prompt"""
        try:
            if not request.text.strip():
                raise ValueError("Text content cannot be empty")

            text = self._sanitize_text(request.text)
            result = await self._analyze_technical_document(text)

            return {"analysis": result}
        except Exception as e:
            self.logger.error(f"Technical document analysis failed: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))

    async def analyze_multi_chunk(self, request: DocumentAnalysisRequest) -> MultiChunkAnalysisResponse:
        """Multi-chunk technical analysis"""
        try:
            if not request.chunks:
                raise ValueError("No chunks provided for analysis")

            analyses = []
            for chunk in request.chunks:
                if not chunk.content.strip():
                    self.logger.warning(f"Empty content in chunk {chunk.chunk_id}")
                    continue

                analysis_result = await self._analyze_chunk_technical(
                    chunk.content,
                    request.user_story
                )

                analyses.append(TechnicalAnalysisResponse(
                    chunk_id=chunk.chunk_id,
                    analysis=analysis_result
                ))

            return MultiChunkAnalysisResponse(
                user_story=request.user_story,
                analyses=analyses
            )
        except Exception as e:
            self.logger.error(f"Multi-chunk analysis failed: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))

    async def analyze_workflow_chunks(self, request: DocumentAnalysisRequest) -> Dict[str, Any]:
        """Analyze chunks following the exact workflow pattern (6 chunks)"""
        try:
            if len(request.chunks) != 6:
                raise HTTPException(
                    status_code=400,
                    detail="This endpoint requires exactly 6 documentation chunks"
                )

            results = {}
            for i, chunk in enumerate(request.chunks, 1):
                chunk_key = f"tech_analysis_chunk_{i}"
                analysis_result = await self._analyze_chunk_technical(
                    chunk.content,
                    request.user_story
                )
                results[chunk_key] = analysis_result

            return {
                "user_story": request.user_story,
                "workflow_version": "0.7",
                "analyses": results
            }
        except Exception as e:
            self.logger.error(f"Workflow chunks analysis failed: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))

    async def _analyze_document(self, text: str) -> str:
        """Original document analysis logic"""
        prompt = f"""
        Please analyze the following document and provide a comprehensive analysis:

        Document Content:
        {text}

        Please provide analysis covering:
        1. Main topics and themes
        2. Technical specifications (if any)
        3. Key requirements or objectives
        4. Potential risks or considerations
        5. Implementation suggestions
        """

        return await self.azure_client.generate_text(
            system_prompt="You are a professional document analyst with expertise in technical documentation.",
            user_prompt=prompt
        )

    async def _analyze_technical_document(self, text: str) -> str:
        """Technical document analysis using workflow system prompt"""
        user_prompt = f"""Analyze this documentation chunk:

        {text}

        Focus on technical details for system understanding."""

        return await self.azure_client.generate_text(
            system_prompt=self.TECHNICAL_ANALYSIS_SYSTEM_PROMPT,
            user_prompt=user_prompt
        )

    async def _analyze_chunk_technical(self, chunk_content: str, user_story: str) -> str:
        """Analyze individual chunk following workflow pattern"""
        user_prompt = f"""User Story: {user_story}

        Documentation Chunk:
        {chunk_content}"""

        return await self.azure_client.generate_text(
            system_prompt=self.TECHNICAL_ANALYSIS_SYSTEM_PROMPT,
            user_prompt=user_prompt
        )


class DocumentAnalysisController:
    def __init__(self, generator: Optional[TextGenerator] = None):
        self.service = DocumentAnalysisService(generator)
        self.logger = logging.getLogger(__name__)
